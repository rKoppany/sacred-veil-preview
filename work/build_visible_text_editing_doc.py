import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PROJECT = Path(r"C:\Users\rolan\Desktop\Astro Website")
WORK = PROJECT / "work"
DATA = json.loads((WORK / "text-style-inventory.json").read_text(encoding="utf-8"))
OUT_DOCX = PROJECT / "Sacred Veil - lathato szovegek modositasi munkatabla.docx"

PAGE_ORDER = [
    "index.html",
    "rolunk.html",
    "szolgaltatasok.html",
    "portfolio.html",
    "kapcsolat.html",
    "aszf.html",
]

PAGE_NAMES = {
    "index.html": "Főoldal",
    "rolunk.html": "Rólunk",
    "szolgaltatasok.html": "Csomagok",
    "portfolio.html": "Portfólió",
    "kapcsolat.html": "Kapcsolat",
    "aszf.html": "ÁSZF",
}

FONT_KEYS = (
    "font-family",
    "font-size",
    "font-weight",
    "font-style",
    "color",
    "letter-spacing",
    "text-transform",
    "text-decoration",
    "--tw-text-opacity",
)

PARAGRAPH_KEYS = (
    "line-height",
    "margin",
    "margin-top",
    "margin-bottom",
    "margin-left",
    "margin-right",
    "padding",
    "padding-top",
    "padding-bottom",
    "padding-left",
    "padding-right",
    "display",
    "gap",
    "row-gap",
    "column-gap",
    "max-width",
    "width",
    "min-height",
    "height",
    "align-items",
    "justify-content",
    "justify-self",
    "text-align",
    "border",
    "border-radius",
    "box-shadow",
    "background",
    "grid-template-columns",
)


def clean_spaces(text):
    return re.sub(r"\s+", " ", text or "").strip()


def shorten(text, limit=760):
    text = clean_spaces(text)
    if len(text) <= limit:
        return text
    return text[: limit - 1].rstrip() + "..."


def rel(path):
    try:
        return str(Path(path).relative_to(PROJECT)).replace("\\", "/")
    except Exception:
        return str(path or "").replace("\\", "/")


def source_ref(matches):
    if not matches:
        return "Forrás: pontos forrássor nem azonosítható automatikusan."
    parts = []
    for match in matches[:4]:
        parts.append(f"{rel(match.get('file'))}:{match.get('line')} ({match.get('match')})")
    return "; ".join(parts)


def is_visible_text_entry(entry):
    kind = entry.get("kind", "")
    tag = entry.get("tag", "")
    classes = entry.get("classes") or ""
    path = entry.get("path") or ""
    text = clean_spaces(entry.get("text"))

    if not text:
        return False
    if tag in {"title", "meta", "script", "style"}:
        return False
    if kind.startswith("aria") or kind.startswith("alt"):
        return False
    if "sr-only" in classes:
        return False
    if "contact-honeypot" in classes or "contact-honeypot" in path:
        return False
    return True


def is_header_entry(entry):
    path = entry.get("path") or ""
    return path.startswith("html > body > header") or " > header." in path


def is_footer_entry(entry):
    path = entry.get("path") or ""
    return path.startswith("html > body > footer") or " > footer." in path


def unique_entries(entries):
    seen = set()
    unique = []
    for entry in entries:
        key = (
            entry.get("text"),
            entry.get("tag"),
            entry.get("classes") or "",
            "header" if is_header_entry(entry) else "footer" if is_footer_entry(entry) else entry.get("path"),
        )
        if key in seen:
            continue
        seen.add(key)
        unique.append(entry)
    return unique


def row_type(entry):
    tag = entry.get("tag") or ""
    kind = entry.get("kind") or ""
    classes = entry.get("classes") or ""
    if kind == "placeholder":
        return "Placeholder / űrlapmezőben látható segédszöveg"
    if tag == "option":
        return "Lenyíló menü opció"
    if tag == "label":
        return "Űrlap felirat"
    if tag in {"button"} or "button" in classes.split():
        return "Gombszöveg"
    if tag == "a":
        return "Link vagy menüpont"
    if tag in {"h1", "h2", "h3", "h4"}:
        return "Cím"
    if tag == "li":
        return "Listaelem"
    return "Látható szöveg"


def style_lines(entry, keys):
    lines = []
    for item in entry.get("style_summary", []):
        low = item.lower()
        if any(key in low for key in keys):
            lines.append("- " + item)
    return lines


def infer_font(entry):
    classes = entry.get("classes") or ""
    container = entry.get("container") or {}
    combined_parts = [
        "\n".join(entry.get("style_summary", [])),
        classes,
        container.get("classes") or "",
    ]
    for rule in container.get("rules", []):
        props = rule.get("props", {})
        if "font-family" in props:
            combined_parts.append(props["font-family"])
        combined_parts.append(rule.get("selector") or "")
    combined = " ".join(combined_parts).lower()
    if (
        "cormorant" in combined
        or "font-display" in combined
        or "display-title" in combined
        or "section-title" in combined
        or "page-title" in combined
    ):
        return "Cormorant Garamond"
    return "Inter"


def html_info(entry):
    classes = entry.get("classes") or "-"
    item_id = entry.get("id") or "-"
    path = entry.get("path") or "-"
    suggestion = entry.get("suggestion") or {}
    parts = [
        f"Oldal: {entry.get('page')} ({entry.get('page_file')})",
        f"id: {item_id}",
        f"class: {shorten(classes, 420)}",
        f"path: {shorten(path, 520)}",
        source_ref(entry.get("source_matches", [])),
    ]
    if entry.get("href"):
        parts.append(f"href: {entry.get('href')}")
    if entry.get("name"):
        parts.append(f"name: {entry.get('name')}")
    if suggestion.get("html"):
        parts.append("Külön kezeléshez HTML class hozzáadás:")
        parts.append(shorten(suggestion["html"], 520))
    return "\n".join(parts)


def paragraph_container_info(entry):
    lines = []
    para = style_lines(entry, PARAGRAPH_KEYS)
    if para:
        lines.append("Közvetlen paragraph/layout szabályok:")
        lines.extend(para[:8])
    else:
        lines.append("Közvetlen paragraph/layout szabály: nincs külön találat; az alap body line-height és a konténer öröklődés érvényesül.")

    container = entry.get("container") or {}
    if container.get("descriptor"):
        lines.append("")
        lines.append("Konténer path:")
        lines.append(shorten(container["descriptor"], 520))
    if container.get("note"):
        lines.append("Irányhatás:")
        lines.append(container["note"])
    rules = []
    for rule in container.get("rules", [])[:5]:
        props = rule.get("props", {})
        picked = []
        for key in PARAGRAPH_KEYS:
            if key in props:
                picked.append(f"{key}: {props[key]}")
        if picked:
            rules.append(f"{rule.get('selector')}: " + "; ".join(picked[:8]))
    if rules:
        lines.append("Konténer CSS:")
        lines.extend("- " + r for r in rules)
    lines.append("")
    lines.append("Szöveg körüli hatás értelmezése:")
    lines.append(direction_summary(entry))
    return "\n".join(lines)


def direction_summary(entry):
    classes = f"{entry.get('classes') or ''} {(entry.get('container') or {}).get('classes') or ''}"
    notes = []
    for cls in classes.split():
        if cls.startswith("mt-"):
            notes.append(f"fent: külső margó ({cls})")
        elif cls.startswith("mb-"):
            notes.append(f"lent: külső margó ({cls})")
        elif cls.startswith("mx-"):
            notes.append(f"bal+jobb: külső margó ({cls})")
        elif cls.startswith("my-"):
            notes.append(f"fent+lent: külső margó ({cls})")
        elif cls.startswith("px-"):
            notes.append(f"bal+jobb: belső padding ({cls})")
        elif cls.startswith("py-"):
            notes.append(f"fent+lent: belső padding ({cls})")
        elif cls.startswith("pt-"):
            notes.append(f"fent: belső padding ({cls})")
        elif cls.startswith("pb-"):
            notes.append(f"lent: belső padding ({cls})")
        elif cls.startswith("pl-"):
            notes.append(f"bal: belső padding ({cls})")
        elif cls.startswith("pr-"):
            notes.append(f"jobb: belső padding ({cls})")
        elif cls.startswith("gap-"):
            notes.append(f"elemek között: rés/távolság ({cls})")
    if not notes:
        return "Nincs egyértelmű irány szerinti spacing class; a távolságot valószínűleg szülő konténer, grid/flex gap vagy alap bekezdésritmus adja."
    return "; ".join(dict.fromkeys(notes))


def css_info(entry):
    lines = []
    font_related = style_lines(entry, FONT_KEYS)
    if font_related:
        lines.append("Font/szín selectorok és paraméterek:")
        lines.extend(font_related[:9])
    else:
        lines.append("Nincs közvetlen font/szín selector; az alap body vagy szülő class öröklődik.")

    suggestion = entry.get("suggestion") or {}
    if suggestion.get("css"):
        lines.append("")
        lines.append("Egyedi módosítás új selectorral:")
        lines.append(suggestion["css"])
    lines.append("")
    lines.append("Tipográfiailag jellemzően módosítani lehet: font-family, font-size, font-weight, color, letter-spacing, text-transform, text-shadow.")
    return "\n".join(lines)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D8D2C8", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_width(cell, cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(cm * 567)))


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def clear_cell(cell):
    for p in cell.paragraphs:
        p.clear()


def add_cell_text(cell, text, size=7.0, bold=False, color="22201D"):
    clear_cell(cell)
    parts = str(text or "").split("\n")
    for i, part in enumerate(parts):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(part)
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = "Heading 1" if level == 1 else "Heading 2"
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(18 if level == 1 else 14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(125, 48, 66)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(80, 74, 68)


def make_table(doc, entries):
    headers = ["#/ típus", "Szöveg", "Font", "Html", "Paragraph/Container", "Css", "Módosítások"]
    widths = [1.6, 4.2, 4.1, 5.0, 5.2, 5.1, 3.1]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, "EDE7DD")
        set_cell_border(cell, "B9A98F", "6")
        set_cell_width(cell, widths[i])
        add_cell_text(cell, h, size=8.2, bold=True, color="22201D")

    for number, entry in enumerate(entries, 1):
        row = table.add_row()
        cells = row.cells
        for i, cell in enumerate(cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_border(cell)
            set_cell_width(cell, widths[i])

        add_cell_text(cells[0], f"{number}\n{row_type(entry)}", size=6.4)
        add_cell_text(cells[1], f"\"{entry.get('text')}\"", size=7.2)
        add_cell_text(cells[2], infer_font(entry), size=6.4)
        add_cell_text(cells[3], html_info(entry), size=6.2)
        add_cell_text(cells[4], paragraph_container_info(entry), size=6.2)
        add_cell_text(cells[5], css_info(entry), size=6.2)
        add_cell_text(cells[6], "", size=7.0)

    doc.add_paragraph()
    return table


def visible_entries(page):
    return [e for e in DATA["entries"][page] if is_visible_text_entry(e)]


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(42.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    styles["Normal"].font.size = Pt(8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Sacred Veil - látható szövegek tipográfiai módosítási munkatábla")
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(34, 32, 29)

    add_note(
        doc,
        "Cél: kizárólag a weboldalon látható szövegek összegyűjtése olyan formában, "
        "hogy a betűtípus, szín, méret, spacing, konténer és CSS azonosítás vizuális "
        "szerkesztéshez is használható legyen. Alt, aria, meta és HTML title szövegek "
        "nem szerepelnek. A Módosítások oszlop szándékosan üres.",
    )
    add_note(doc, f"Forrásprojekt: {PROJECT} | Generálva: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    index_entries = visible_entries("index.html")
    nav_entries = unique_entries([e for e in index_entries if is_header_entry(e)])
    footer_entries = unique_entries([e for e in index_entries if is_footer_entry(e)])

    add_heading(doc, "Navigációs sáv - minden oldalon közös", level=1)
    add_note(doc, "A fejléc közös elemei: Sacred Veil felirat, menüpontok, ajánlatkérő gomb, valamint mobilnézetben a Menü gomb és a mobil menüpontok.")
    make_table(doc, nav_entries)

    add_heading(doc, "Footer - minden oldalon közös", level=1)
    add_note(doc, "A lábléc közös elemei egyesével. Ezek az oldalspecifikus táblázatokban már nem ismétlődnek.")
    make_table(doc, footer_entries)

    for page in PAGE_ORDER:
        entries = [
            e
            for e in visible_entries(page)
            if not is_header_entry(e) and not is_footer_entry(e)
        ]
        add_heading(doc, f"{PAGE_NAMES[page]} ({page})", level=1)
        add_note(doc, f"Látható, csak ezen az oldalon szereplő szöveges elemek száma: {len(entries)}.")
        make_table(doc, entries)

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_doc()
