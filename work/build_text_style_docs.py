import json
import re
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

PROJECT = Path(r"C:\Users\rolan\Desktop\Astro Website")
WORK = PROJECT / "work"
DATA = json.loads((WORK / "text-style-inventory.json").read_text(encoding="utf-8"))
OUT_DOCX = PROJECT / "Sacred Veil - szovegek es tipografia dokumentacio.docx"
OUT_PDF = PROJECT / "Sacred Veil - szovegek es tipografia dokumentacio.pdf"

PAGE_ORDER = ["index.html", "rolunk.html", "szolgaltatasok.html", "portfolio.html", "kapcsolat.html", "aszf.html"]
PAGE_NAMES = DATA["pages"]

# ---------- helpers ----------
def rel(path):
    if not path:
        return "-"
    try:
        return str(Path(path).relative_to(PROJECT)).replace("\\", "/")
    except Exception:
        return str(path).replace("\\", "/")

def esc(s):
    return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

def shorten(s, n=420):
    s = re.sub(r"\s+", " ", s or "").strip()
    return s if len(s) <= n else s[: n - 1].rstrip() + "…"

def source_ref(matches):
    if not matches:
        return "Forráskeresés: nincs pontos találat; valószínűleg generált vagy attribútumból képzett szöveg."
    parts = []
    for m in matches:
        parts.append(f"{rel(m['file'])}:{m['line']} ({m['match']})")
    return "; ".join(parts)

def style_short(entry):
    bits = entry.get("style_summary", [])[:5]
    return "\n".join(f"- {shorten(b, 240)}" for b in bits) or "- Örökölt alapstílus."

def container_short(entry):
    c = entry.get("container", {})
    rules = []
    for r in c.get("rules", [])[:3]:
        props = r.get("props", {})
        keep = []
        for k in ["padding", "padding-top", "padding-bottom", "padding-left", "padding-right", "margin-top", "margin-bottom", "display", "gap", "text-align", "max-width", "background"]:
            if k in props:
                keep.append(f"{k}: {props[k]}")
        if keep:
            rules.append(f"{r['selector']}: " + "; ".join(keep[:4]))
    out = [f"Konténer: {shorten(c.get('descriptor','-'), 260)}"]
    if c.get("note"):
        out.append(f"Irányhatás: {shorten(c['note'], 320)}")
    if rules:
        out.append("CSS: " + " | ".join(rules))
    return "\n".join(out)

def modification_short(entry):
    sg = entry.get("suggestion", {})
    html = sg.get("html", "")
    css = sg.get("css", "")
    return f"Külön kezeléshez HTML-ben class hozzáadása:\n{html}\nCSS selector:\n{css}"

def is_header_footer(entry):
    path = entry.get("path", "")
    return path.startswith("header") or path.startswith("footer") or " > header" in path or " > footer" in path

def is_portfolio_generated(entry):
    if entry.get("page_file") != "portfolio.html":
        return False
    return bool(re.fullmatch(r"Sacred Veil portfóliófotó \d+( megnyitása)?", entry.get("text", "")))

def grouped_page_entries(file):
    entries = DATA["entries"][file]
    result = []
    # Page meta pseudo entries
    meta = DATA["meta"][file]
    result.append({
        "index": 0,
        "page": PAGE_NAMES[file], "page_file": file, "kind": "HTML title", "text": meta.get("title", ""),
        "tag": "title", "classes": "", "id": "", "path": "head > title", "style_origin": "Layout.astro title prop",
        "style_summary": ["Böngészőfül/SEO cím, nem oldaltipográfia."],
        "container": {"descriptor": "head", "classes": "", "rules": [], "note": "Nincs vizuális spacing."},
        "source_matches": [],
        "suggestion": {"html": '<Layout title="...">', "css": "Nem CSS-ben módosítandó."}
    })
    result.append({
        "index": 0,
        "page": PAGE_NAMES[file], "page_file": file, "kind": "meta description", "text": meta.get("description", ""),
        "tag": "meta", "classes": "", "id": "", "path": "head > meta[name=description]", "style_origin": "Layout.astro description prop",
        "style_summary": ["SEO leírás, nem oldaltipográfia."],
        "container": {"descriptor": "head", "classes": "", "rules": [], "note": "Nincs vizuális spacing."},
        "source_matches": [],
        "suggestion": {"html": '<Layout description="...">', "css": "Nem CSS-ben módosítandó."}
    })
    if file == "portfolio.html":
        added_alt = False
        added_aria = False
        for e in entries:
            if is_header_footer(e):
                continue
            t = e["text"]
            if re.fullmatch(r"Sacred Veil portfóliófotó \d+", t):
                if not added_alt:
                    ee = dict(e)
                    ee["text"] = "Sacred Veil portfóliófotó 1-144"
                    ee["kind"] = "generált alt szöveg tartomány"
                    ee["source_matches"] = [{"file": str(PROJECT / "src/components/PortfolioGrid.astro"), "line": 94, "match": "template"}]
                    result.append(ee); added_alt = True
                continue
            if re.fullmatch(r"Sacred Veil portfóliófotó \d+ megnyitása", t):
                if not added_aria:
                    ee = dict(e)
                    ee["text"] = "Sacred Veil portfóliófotó 1-144 megnyitása"
                    ee["kind"] = "generált aria-label tartomány"
                    ee["source_matches"] = [{"file": str(PROJECT / "src/components/PortfolioGrid.astro"), "line": 91, "match": "template"}]
                    result.append(ee); added_aria = True
                continue
            result.append(e)
    else:
        result.extend([e for e in entries if not is_header_footer(e)])
    return result

shared_entries = []
seen = set()
for e in DATA["entries"]["index.html"]:
    if is_header_footer(e):
        key = (e["kind"], e["text"], e["tag"], e.get("classes", ""), e.get("path", "").split(" > ")[0])
        if key not in seen:
            ee = dict(e)
            ee["page"] = "Minden oldal (Header/Footer)"
            shared_entries.append(ee)
            seen.add(key)

# ---------- DOCX styles ----------
def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def set_cell_border(cell, color="DADCE0", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), sz)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)

def add_para(cell, text, bold_prefix=False, size=8):
    p = cell.paragraphs[0] if not cell.paragraphs[0].text else cell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    r.font.size = Pt(size)
    if bold_prefix:
        r.bold = True
    return p

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    return p

def add_note(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    set_cell_border(cell, "DADCE0", "4")
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    r = cell.paragraphs[0].add_run(text)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(67, 67, 67)
    doc.add_paragraph()

def make_docx():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.72)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(0.65)
    sec.right_margin = Inches(0.65)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color in [("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 11, "1F4D78")]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(10)
        st.paragraph_format.space_after = Pt(5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Sacred Veil weboldal - szövegek és tipográfiai azonosítók")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string("22201D")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Forrásprojekt: {PROJECT}\nGenerálva: {datetime.now().strftime('%Y-%m-%d %H:%M')}").font.size = Pt(9)
    add_note(doc, "A dokumentum célja: minden fontos oldalszöveg, menüpont, gombszöveg, űrlapszöveg, SEO cím/leírás, alt/aria szöveg, valamint a hozzájuk kapcsolódó tipográfiai és konténer-azonosítók kereshető leltára. A legtöbb stílus HTML classként van megadva Astro fájlokban, és a Tailwind build vagy a src/styles/global.css fordítja tényleges CSS-é.")

    add_heading(doc, "1. Projekt- és fájlstruktúra", 1)
    rows = [
        ["Fájl", "Szerep", "Szöveg/stílus módosításának jellege"],
        ["src/components/Layout.astro", "Közös HTML fej, Header, Footer, main, Google Fonts, veil-transition script.", "SEO title/description propok itt kerülnek head-be; globális script link itt található."],
        ["src/components/Header.astro", "Felső navigáció, mobil menü, logó, fő CTA.", "Menüpontok a src/data/site.ts navItems tömbből jönnek; gombszöveg közvetlenül itt."],
        ["src/components/Footer.astro", "Lábléc márkaszöveg, oldallista, kapcsolatblokk, copyright.", "Ismétlődő footer szövegek itt módosíthatók."],
        ["src/components/PageHero.astro", "Általános aloldali hero komponens.", "Eyebrow/title/lead propként érkezik az oldal Astro fájljából; classok itt vannak."],
        ["src/components/PortfolioGrid.astro", "Portfólió képrács, lightbox, portfólió animációk.", "Képekből generált alt/aria szövegek és lightbox vezérlők itt vannak; inline style blokkja buildkor CSS-be kerül."],
        ["src/data/site.ts", "Navigáció, kiemelt portfólióképek, szolgáltatás-adatok.", "Menüpont-címkék, portfólió alt minták, szolgáltatáscsoportok szövegei."],
        ["src/pages/*.astro", "Oldalankénti tartalom.", "Az oldalspecifikus szövegek döntő része itt módosítható."],
        ["src/styles/global.css", "Globális saját CSS és Tailwind layer szabályok.", "Közös tipográfia: body, .eyebrow, .display-title, .page-title, .section-title, .button, package/portfolio/lightbox stb."],
        ["tailwind.config.mjs", "Szín- és font tokenek.", "font-display, font-body, ivory/veil/champagne/rosegold/taupe/ink színek."],
    ]
    t = doc.add_table(rows=len(rows), cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            cell = t.cell(i, j)
            if i == 0:
                set_cell_shading(cell, "E8EEF5")
            set_cell_border(cell)
            add_para(cell, text, size=8.5 if i else 9)
            if i == 0:
                cell.paragraphs[0].runs[0].bold = True
    set_repeat_table_header(t.rows[0])

    add_heading(doc, "2. Közös tipográfiai rendszer", 1)
    add_note(doc, "A tényleges webes betűkészlet: Inter a body/alapszöveg, Cormorant Garamond a display jellegű címeknél. A classok nagy része Tailwind utility, ezért sok tulajdonság a HTML class listájában kereshető, nem kézzel írt CSS blokkban.")
    rows = [["Selector / token", "Fő tulajdonságok", "Hol kereshető"]]
    for item in DATA["shared_styles"]:
        props = item["props"]
        prop_text = "; ".join(f"{k}: {v}" for k, v in props.items() if k in {"font-family","font-size","font-weight","line-height","letter-spacing","text-transform","color","padding","margin-inline","width","background","border-radius","box-shadow","display","gap"})
        rows.append([item["selector"], shorten(prop_text, 620), "src/styles/global.css"])
    rows += [
        ["font-display", "Cormorant Garamond, Georgia, serif", "tailwind.config.mjs -> theme.extend.fontFamily.display"],
        ["font-body/body", "Inter, Segoe UI, sans-serif", "tailwind.config.mjs + body szabály a global.css-ben"],
        ["text-ink / bg-ink", "#22201d", "tailwind.config.mjs -> colors.ink"],
        ["text-taupe", "#6f6860", "tailwind.config.mjs -> colors.taupe"],
        ["text-champagne", "#b08a4d", "tailwind.config.mjs -> colors.champagne"],
        ["text-rosegold", "#7d3042", "tailwind.config.mjs -> colors.rosegold"],
    ]
    t = doc.add_table(rows=len(rows), cols=3)
    t.style = "Table Grid"
    widths = [1.6, 4.2, 1.7]
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            cell = t.cell(i, j)
            if i == 0:
                set_cell_shading(cell, "E8EEF5")
            set_cell_border(cell)
            add_para(cell, text, size=8)
            if i == 0: cell.paragraphs[0].runs[0].bold = True
    set_repeat_table_header(t.rows[0])

    add_heading(doc, "3. Külön formázás általános menete", 1)
    steps = [
        "Az adott szöveg forrássorának megkeresése a dokumentumban szereplő fájl:line jelöléssel vagy a beidézett szöveg Ctrl+F keresésével.",
        "Ha a szöveg közös classon van, például .section-title vagy .button, akkor külön class hozzáadása ajánlott az adott HTML taghez.",
        "A hozzáadott class egyedi selectorral formázható a src/styles/global.css végén vagy egy logikailag megfelelő szakaszban.",
        "Ha a stílus Tailwind utility classból jön, akkor gyors módosításhoz a classlista változtatható az Astro fájlban; hosszabb távon az egyedi class tisztább.",
        "Ha a szöveg adatból generált, például navItems, packages vagy portfolioImages, akkor a szövegforrás tömbjét kell módosítani, nem a dist HTML-t.",
    ]
    for s in steps:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(s)

    add_heading(doc, "4. Ismétlődő Header/Footer szövegek", 1)
    add_entries_table(doc, shared_entries, compact=True)

    add_heading(doc, "5. Oldalankénti szövegkatalógus", 1)
    for file in PAGE_ORDER:
        doc.add_page_break()
        add_heading(doc, f"{PAGE_NAMES[file]} ({file})", 1)
        meta = DATA["meta"][file]
        add_note(doc, f"HTML title: {meta.get('title','')}\nMeta description: {meta.get('description','')}")
        entries = grouped_page_entries(file)
        add_entries_table(doc, entries, compact=False)

    doc.add_page_break()
    add_heading(doc, "6. Megjegyzések a nem konvencionális működéshez", 1)
    notes = [
        ("Egy globális CSS fájl", "A projekt Astro + Tailwind felépítésű, ezért a src/styles/global.css egyszer importálódik a Layout.astro-ban. Ez biztosítja, hogy a fátyolátmenet, a közös komponensek és az oldalanként ismétlődő elemek ugyanazt a stílusrendszert használják."),
        ("Tailwind utility classok", "Sok tipográfiai beállítás nem külön CSS blokkban van, hanem HTML classként: például text-lg, font-semibold, text-taupe/78, mt-6. Ezeket az Astro fájlokban kell keresni."),
        ("Fátyol transition JS", "A public/js/veil-transition.js nem csak navigációs animációt végez, hanem screenshot/canvas rétegeket, lace mask feldolgozást és felfedő maszkot is kezel. Emiatt a hozzá tartozó CSS osztályoknak globálisnak kell maradniuk."),
        ("Portfólió inline komponenslogika", "A PortfolioGrid.astro buildkor generálja a képek sorrendjét, méreteit, alt/aria szövegét és a lightbox működését. A portfólió szövegeinek egy része ezért template-stringből keletkezik."),
        ("Közös class külön kezelése", "Ha egy .section-title stílusú cím csak egyetlen helyen legyen más, nem a .section-title globális szabályt érdemes átírni, hanem az adott h2 kapjon plusz egyedi class-t, például home-closing-title."),
    ]
    for title_text, body in notes:
        p = doc.add_paragraph()
        p.style = "Heading 2"
        p.add_run(title_text)
        doc.add_paragraph(body)

    doc.save(OUT_DOCX)


def add_entries_table(doc, entries, compact=False):
    headers = ["# / típus", "Szövegblokk", "Azonosítás és forrás", "Stílus és tipográfia", "Konténer / külön formázás"]
    t = doc.add_table(rows=1, cols=5)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, "E8EEF5")
        set_cell_border(cell)
        add_para(cell, h, size=7.6)
        cell.paragraphs[0].runs[0].bold = True
    limit = None
    for n, e in enumerate(entries, 1):
        row = t.add_row()
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_border(cell, "DADCE0", "3")
        idx = e.get("index") or n
        add_para(row.cells[0], f"{idx}\n{e.get('kind','')}", size=7.2)
        add_para(row.cells[1], f"„{shorten(e.get('text',''), 900 if not compact else 420)}”", size=7.2)
        ident = [
            f"Tag: <{e.get('tag','')}>",
            f"id: {e.get('id') or '-'}",
            f"class: {shorten(e.get('classes') or '-', 340)}",
            f"path: {shorten(e.get('path') or '-', 360)}",
            f"forrás: {source_ref(e.get('source_matches', []))}",
            f"stílus helye: {e.get('style_origin','-')}",
        ]
        if e.get("href"):
            ident.append(f"href: {e['href']}")
        if e.get("name"):
            ident.append(f"name: {e['name']}")
        add_para(row.cells[2], "\n".join(ident), size=7.0)
        add_para(row.cells[3], style_short(e), size=7.0)
        add_para(row.cells[4], container_short(e) + "\n\n" + modification_short(e), size=6.8)
    # Width hints via tcW
    dxa_widths = [720, 2280, 2340, 2500, 2700]
    for row in t.rows:
        for cell, w in zip(row.cells, dxa_widths):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.first_child_found_in("w:tcW")
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(w))
            tcW.set(qn("w:type"), "dxa")
    doc.add_paragraph()

# ---------- PDF ----------
def register_pdf_font():
    candidates = [
        Path(r"C:\Windows\Fonts\calibri.ttf"),
        Path(r"C:\Windows\Fonts\arial.ttf"),
        Path(r"C:\Windows\Fonts\segoeui.ttf"),
    ]
    for p in candidates:
        if p.exists():
            pdfmetrics.registerFont(TTFont("DocFont", str(p)))
            b = p.with_name(p.stem + "bd" + p.suffix)
            if b.exists():
                pdfmetrics.registerFont(TTFont("DocFont-Bold", str(b)))
            else:
                pdfmetrics.registerFont(TTFont("DocFont-Bold", str(p)))
            return
    # fallback built-in

PDF_FONT = "DocFont"
PDF_BOLD = "DocFont-Bold"

def para(text, style):
    return Paragraph(esc(text).replace("\n", "<br/>"), style)

def make_pdf():
    register_pdf_font()
    styles = getSampleStyleSheet()
    body = ParagraphStyle("BodySV", parent=styles["BodyText"], fontName=PDF_FONT, fontSize=8.2, leading=10.2, spaceAfter=4)
    small = ParagraphStyle("SmallSV", parent=body, fontSize=6.6, leading=8.0)
    title = ParagraphStyle("TitleSV", parent=styles["Title"], fontName=PDF_BOLD, fontSize=18, leading=22, alignment=TA_CENTER, textColor=colors.HexColor("#22201D"), spaceAfter=12)
    h1 = ParagraphStyle("H1SV", parent=styles["Heading1"], fontName=PDF_BOLD, fontSize=14, leading=17, textColor=colors.HexColor("#2E74B5"), spaceBefore=12, spaceAfter=6)
    h2 = ParagraphStyle("H2SV", parent=styles["Heading2"], fontName=PDF_BOLD, fontSize=11, leading=13, textColor=colors.HexColor("#1F4D78"), spaceBefore=8, spaceAfter=5)
    note = ParagraphStyle("NoteSV", parent=body, fontSize=8, leading=10, backColor=colors.HexColor("#F4F6F9"), borderColor=colors.HexColor("#DADCE0"), borderWidth=0.4, borderPadding=5, spaceAfter=8)
    doc = SimpleDocTemplate(str(OUT_PDF), pagesize=A4, rightMargin=1.1*cm, leftMargin=1.1*cm, topMargin=1.1*cm, bottomMargin=1.1*cm)
    story = []
    story.append(Paragraph("Sacred Veil weboldal - szövegek és tipográfiai azonosítók", title))
    story.append(Paragraph(f"Forrásprojekt: {esc(str(PROJECT))}<br/>Generálva: {datetime.now().strftime('%Y-%m-%d %H:%M')}", body))
    story.append(Paragraph("A dokumentum kereshető referencia a weboldal szövegeihez, class/id azonosítóihoz, tipográfiai forrásaihoz és külön formázási lehetőségeihez.", note))
    story.append(Paragraph("1. Projekt- és fájlstruktúra", h1))
    rows = [[para("Fájl", small), para("Szerep", small), para("Módosítási irány", small)]]
    for r in [
        ["src/components/Layout.astro", "Közös HTML fej, Header, Footer, main, Google Fonts, transition script.", "SEO propok, közös script link."],
        ["src/components/Header.astro", "Felső navigáció és mobil menü.", "Menüpontok: src/data/site.ts; CTA közvetlenül itt."],
        ["src/components/Footer.astro", "Lábléc márkaszöveg, oldallista, kapcsolat.", "Footer szövegek itt."],
        ["src/components/PageHero.astro", "Aloldali hero komponens.", "Eyebrow/title/lead propok oldalból."],
        ["src/components/PortfolioGrid.astro", "Portfólió képrács és lightbox.", "Generált alt/aria, lightbox gombok, JS/CSS."],
        ["src/styles/global.css", "Globális saját CSS.", "Közös tipográfia és komponensstílusok."],
        ["tailwind.config.mjs", "Font/szín tokenek.", "font-display, font-body, színek."],
    ]:
        rows.append([para(x, small) for x in r])
    story.append(make_pdf_table(rows, [4.0*cm, 6.0*cm, 8.0*cm]))
    story.append(Paragraph("2. Közös tipográfiai rendszer", h1))
    rows = [[para("Selector / token", small), para("Fő tulajdonságok", small), para("Hol kereshető", small)]]
    for item in DATA["shared_styles"]:
        props = item["props"]
        prop_text = "; ".join(f"{k}: {v}" for k, v in props.items() if k in {"font-family","font-size","font-weight","line-height","letter-spacing","text-transform","color","padding","margin-inline","width","background","border-radius","box-shadow","display","gap"})
        rows.append([para(item["selector"], small), para(shorten(prop_text, 500), small), para("src/styles/global.css", small)])
    rows.extend([
        [para("font-display", small), para("Cormorant Garamond, Georgia, serif", small), para("tailwind.config.mjs", small)],
        [para("body", small), para("Inter, Segoe UI, sans-serif; line-height: 1.65; color: var(--taupe)", small), para("src/styles/global.css", small)],
    ])
    story.append(make_pdf_table(rows, [3.8*cm, 9.0*cm, 5.2*cm]))
    story.append(Paragraph("3. Ismétlődő Header/Footer szövegek", h1))
    story.extend(pdf_entries(shared_entries, small, h2))
    story.append(PageBreak())
    story.append(Paragraph("4. Oldalankénti szövegkatalógus", h1))
    for file in PAGE_ORDER:
        story.append(PageBreak())
        story.append(Paragraph(f"{PAGE_NAMES[file]} ({file})", h1))
        meta = DATA["meta"][file]
        story.append(Paragraph(f"HTML title: {esc(meta.get('title',''))}<br/>Meta description: {esc(meta.get('description',''))}", note))
        story.extend(pdf_entries(grouped_page_entries(file), small, h2))
    story.append(PageBreak())
    story.append(Paragraph("5. Megjegyzések a nem konvencionális működéshez", h1))
    for t, b in [
        ("Egy globális CSS fájl", "A Layout.astro egyszer importálja a src/styles/global.css-t, ezért a közös komponensek és a fátyolátmenet minden oldalon ugyanazt a stílusrendszert kapja."),
        ("Tailwind utility classok", "Sok beállítás HTML classként található: text-lg, font-semibold, mt-6, text-taupe/78. Ezeket az Astro fájlok class attribútumában lehet keresni."),
        ("Fátyol transition JS", "A public/js/veil-transition.js screenshot/canvas rétegeket, lace mask feldolgozást és felfedő maszkot kezel; ezért globális classokat használ."),
        ("Portfólió generálás", "A PortfolioGrid.astro a képekből generálja a sorrendet, alt/aria szövegeket és lightbox működést."),
    ]:
        story.append(Paragraph(t, h2)); story.append(Paragraph(b, body))
    doc.build(story)

def make_pdf_table(rows, col_widths):
    t = Table(rows, colWidths=col_widths, repeatRows=1)
    t.setStyle(TableStyle([
        ("GRID", (0,0), (-1,-1), 0.35, colors.HexColor("#DADCE0")),
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#E8EEF5")),
        ("VALIGN", (0,0), (-1,-1), "TOP"),
        ("LEFTPADDING", (0,0), (-1,-1), 4),
        ("RIGHTPADDING", (0,0), (-1,-1), 4),
        ("TOPPADDING", (0,0), (-1,-1), 3),
        ("BOTTOMPADDING", (0,0), (-1,-1), 3),
    ]))
    return t

def pdf_entries(entries, small, h2):
    story = []
    rows = [[para("# / típus", small), para("Szöveg", small), para("Azonosítás / forrás", small), para("Stílus / konténer / külön kezelés", small)]]
    for n, e in enumerate(entries, 1):
        idx = e.get("index") or n
        ident = "\n".join([
            f"Tag: <{e.get('tag','')}>",
            f"id: {e.get('id') or '-'}",
            f"class: {shorten(e.get('classes') or '-', 240)}",
            f"forrás: {source_ref(e.get('source_matches', []))}",
        ])
        sty = style_short(e) + "\n\n" + container_short(e) + "\n\n" + modification_short(e)
        rows.append([para(f"{idx}\n{e.get('kind','')}", small), para("„" + shorten(e.get('text',''), 600) + "”", small), para(ident, small), para(sty, small)])
        if len(rows) >= 28:
            story.append(make_pdf_table(rows, [1.7*cm, 5.0*cm, 5.5*cm, 6.2*cm])); story.append(Spacer(1, 7)); rows = rows[:1]
    if len(rows) > 1:
        story.append(make_pdf_table(rows, [1.7*cm, 5.0*cm, 5.5*cm, 6.2*cm])); story.append(Spacer(1, 7))
    return story

make_docx()
make_pdf()
print(OUT_DOCX)
print(OUT_PDF)
